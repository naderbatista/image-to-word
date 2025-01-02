import pytesseract
from PIL import Image
import os
import time
from docx import Document
from docx.shared import Pt

# CONFIGURAR O CAMINHO DO EXECUTÁVEL DO TESSERACT (CASO SEJA NECESSÁRIO)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text_from_image(image_path):
    """EXTRAI TEXTO DE UMA IMAGEM USANDO TESSERACT OCR."""
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"ARQUIVO NÃO ENCONTRADO: {image_path}")
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image, lang='eng')  # 'POR' PARA PORTUGUÊS OU 'ENG' PARA INGLÊS
    return text

def save_text_to_doc(text, output_path):
    """SALVA O TEXTO EXTRAÍDO EM UM ARQUIVO .DOCX."""
    doc = Document()
    
    # ADICIONAR O TEXTO EXTRAÍDO COM A FORMATAÇÃO DESEJADA
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # SALVAR O DOCUMENTO
    doc.save(output_path)

def process_images_in_folder(input_folder, output_folder, interval=5):
    """
    PROCESSA TODAS AS IMAGENS NA PASTA ESPECIFICADA.
    CRIA UM ARQUIVO .DOCX PARA CADA IMAGEM COM O MESMO NOME DA IMAGEM.
    """
    if not os.path.exists(input_folder):
        raise FileNotFoundError(f"PASTA NÃO ENCONTRADA: {input_folder}")
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)  # CRIA A PASTA DE SAÍDA SE NÃO EXISTIR
    
    # LISTAR TODOS OS ARQUIVOS DE IMAGEM NA PASTA
    files = [f for f in os.listdir(input_folder) if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    
    if not files:
        print("NENHUMA IMAGEM ENCONTRADA NA PASTA.")
        return

    for file in files:
        image_path = os.path.join(input_folder, file)
        doc_filename = os.path.splitext(file)[0] + ".docx"
        doc_output_path = os.path.join(output_folder, doc_filename)
        
        print(f"LENDO IMAGEM: {file}")
        extracted_text = extract_text_from_image(image_path)
        
        print(f"SALVANDO TEXTO EXTRAÍDO EM: {doc_filename}")
        save_text_to_doc(extracted_text, doc_output_path)
        
        print(f"PROCESSAMENTO CONCLUÍDO PARA: {file}")
        print(f"AGUARDANDO {interval} SEGUNDOS ANTES DE PROCESSAR A PRÓXIMA IMAGEM...\n")
        time.sleep(interval)

def main():
    # CAMINHO DA PASTA DE ENTRADA (IMAGENS) E SAÍDA (RESULTADOS)
    input_folder = r"C:\caminho-do-arquivo\imagem-ocr\imagens"
    output_folder = r"C:\caminho-do-arquivo\imagem-ocr\resultado"
    
    # INTERVALO ENTRE PROCESSAMENTOS (EM SEGUNDOS)
    interval = 5  # AJUSTE CONFORME NECESSÁRIO
    
    # PROCESSAR AS IMAGENS NA PASTA
    process_images_in_folder(input_folder, output_folder, interval)

if __name__ == "__main__":
    main()
